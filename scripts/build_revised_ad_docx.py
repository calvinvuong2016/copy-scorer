"""
Build a Word document of the revised ad with edited portions highlighted.
Reads original_ad.txt and the revised ad from the project root.
Output: Revised Ad - Score 90+ (Highlighted Edits).docx
"""
import re
from pathlib import Path

# Try to use python-docx; if not installed, we'll fail with a clear message
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_COLOR_INDEX
except ImportError:
    raise SystemExit("Install python-docx: pip install python-docx")

# Paths
SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
ORIGINAL_PATH = SCRIPT_DIR / "original_ad.txt"
REVISED_PATH = PROJECT_ROOT / "Revised Ad - Score 90+ (Edits Applied).txt"
OUTPUT_PATH = PROJECT_ROOT / "Revised Ad - Score 90+ (Highlighted Edits).docx"


def extract_revised_body(content: str) -> str:
    """Extract the revised ad body between the two section markers."""
    start_marker = "REVISED AD COPY (FULL)"
    end_marker = "END OF REVISED AD"
    start = content.find(start_marker)
    end = content.find(end_marker)
    if start == -1 or end == -1:
        return content  # fallback: use whole file
    start = content.find("\n", start) + 1
    return content[start:end].strip()


def get_highlight_ranges(original: str, revised: str):
    """Use difflib to find character ranges in revised that are new or changed. Returns set of (start, end) ranges (inclusive)."""
    import difflib
    matcher = difflib.SequenceMatcher(None, original, revised)
    ranges = set()
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ("replace", "insert"):
            # j1:j2 are the indices in revised that are new/changed
            ranges.add((j1, j2))
    return ranges


def merge_ranges(ranges):
    """Merge overlapping or adjacent ranges."""
    if not ranges:
        return []
    sorted_ranges = sorted(ranges)
    merged = [list(sorted_ranges[0])]
    for start, end in sorted_ranges[1:]:
        if start <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], end)
        else:
            merged.append([start, end])
    return merged


def char_in_ranges(i, merged_ranges):
    """True if character index i is inside any merged range."""
    for start, end in merged_ranges:
        if start <= i < end:
            return True
    return False


def main():
    original_text = ORIGINAL_PATH.read_text(encoding="utf-8")
    revised_content = REVISED_PATH.read_text(encoding="utf-8")
    revised_text = extract_revised_body(revised_content)

    ranges = get_highlight_ranges(original_text, revised_text)
    merged = merge_ranges(ranges)

    # Build document
    doc = Document()
    doc.add_heading("Revised Ad Copy — Edits Highlighted", level=0)
    doc.add_paragraph(
        "Yellow highlighting shows text that was added or changed from the original. "
        "Unhighlighted text is unchanged."
    )
    doc.add_paragraph()
    doc.add_heading("What was updated (for 90+ scores)", level=1)
    summary = [
        "• Mechanism: Clear 'why it works' line (ingredients the dermatology office recommends) and sharper offer clarity.",
        "• Bullets: 5 bullets with benefit + mechanism + intrigue added before the CTA (Makepeace/Bencivenga).",
        "• Risk reversal: 'No subscription. No auto-ship. No cost. Just the list.'",
        "• P.S. that sells: P.S. and P.P.S. added to restate proof, benefit, and light urgency (Halbert).",
        "• Offer clarity: One sentence stating exactly what they get; free; no signup.",
        "• Light urgency: 'I'm not sure how long they keep the page up.'",
        "• Meta: ALL CAPS reduced to one emphasis ('at me'); rest lowercase for organic feel.",
        "• Nina: 'Has for 11 years' added to differentiate the discovery moment.",
    ]
    for line in summary:
        doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()
    doc.add_heading("Revised ad (highlighted = edited)", level=1)
    # Body: split revised into paragraphs (double newline or single)
    paragraphs = re.split(r"\n\n+", revised_text)
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        # For this paragraph, determine which characters are highlighted
        # We need to find this paragraph's start index in revised_text
        start_in_revised = revised_text.find(para_text)
        if start_in_revised == -1:
            p.add_run(para_text)
            continue
        end_in_revised = start_in_revised + len(para_text)
        # Build runs: alternate normal / highlighted
        i = 0
        while i < len(para_text):
            is_highlight = char_in_ranges(start_in_revised + i, merged)
            j = i
            while j < len(para_text) and char_in_ranges(start_in_revised + j, merged) == is_highlight:
                j += 1
            run_text = para_text[i:j]
            if run_text:
                run = p.add_run(run_text)
                if is_highlight:
                    run.font.highlight = WD_COLOR_INDEX.YELLOW
            i = j
    doc.add_paragraph()
    doc.add_paragraph("— End of revised ad —")
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
